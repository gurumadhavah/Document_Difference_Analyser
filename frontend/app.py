import streamlit as st
import requests
import streamlit.components.v1 as components
import io
import pypdf
import docx
from docx import Document
import html as html_lib

# 1. Page Config
st.set_page_config(page_title="DocDiff AI", layout="wide")

# --- STATE MANAGEMENT ---
if "dark_mode" not in st.session_state:
    st.session_state.dark_mode = True
if "analysis_complete" not in st.session_state:
    st.session_state.analysis_complete = False
if "api_data" not in st.session_state:
    st.session_state.api_data = None

def toggle_theme():
    st.session_state.dark_mode = not st.session_state.dark_mode

# Define Color Palettes
if st.session_state.dark_mode:
    theme = {
        "bg_color": "#0E1117",
        "text_color": "#FAFAFA",
        "card_bg": "#262730",
        "input_bg": "#1E1E1E",
        "uploader_bg": "#363945",
        "uploader_text": "#FAFAFA",
        "border_color": "#4A4A4A",
        "shadow": "rgba(0,0,0,0.3)",
        "diff_bg": "#1d1e24",
        "diff_text": "#e0e0e0",
        "diff_del_bg": "#442a2d",
        "diff_del_text": "#ffcdd2",
        "diff_add_bg": "#1b3b24",
        "diff_add_text": "#c8e6c9",
        "overlay_bg": "rgba(14, 17, 23, 0.85)"
    }
else:
    theme = {
        "bg_color": "#F4F6F9",
        "text_color": "#333333",
        "card_bg": "#FFFFFF",
        "input_bg": "#FFFFFF",
        "uploader_bg": "#F0F2F6",
        "uploader_text": "#333333",
        "border_color": "#D1D5DB",
        "shadow": "0 2px 5px rgba(0,0,0,0.05)",
        "diff_bg": "#FFFFFF",
        "diff_text": "#000000",
        "diff_del_bg": "#ffebee",
        "diff_del_text": "#c62828",
        "diff_add_bg": "#e8f5e9",
        "diff_add_text": "#2e7d32",
        "overlay_bg": "rgba(255, 255, 255, 0.85)"
    }

# 2. Dynamic CSS Injection
st.markdown(f"""
<style>
    * {{ transition: background-color 0.3s ease, color 0.3s ease; }}

    .stApp {{
        background-color: {theme['bg_color']} !important;
        color: {theme['text_color']} !important;
    }}
    
    h1, h2, h3, h4, h5, h6, p, label, li {{ color: {theme['text_color']} !important; }}
    
    .block-container {{ padding-top: 2rem !important; padding-bottom: 5rem; }}

    /* MAIN CARDS */
    [data-testid="stVerticalBlockBorderWrapper"] {{
        background-color: {theme['card_bg']};
        border: 1px solid {theme['border_color']} !important;
        box-shadow: {theme['shadow']};
        border-radius: 10px;
        padding: 10px;
    }}

    /* --- EXPANDER HEADER FIX --- */
    /* This targets the clickable header part of the expander */
    div[data-testid="stExpander"] details > summary {{
        background-color: {theme['card_bg']} !important;
        color: {theme['text_color']} !important;
        border: 1px solid {theme['border_color']} !important;
        border-radius: 8px;
        transition: all 0.2s ease;
    }}
    
    /* Hover effect for expander */
    div[data-testid="stExpander"] details > summary:hover {{
        border-color: #F39C12 !important;
        color: #F39C12 !important;
    }}
    
    /* Ensure svg icon inside header matches text color */
    div[data-testid="stExpander"] details > summary svg {{
        color: inherit !important;
        fill: inherit !important;
    }}

    /* --- TEXT AREA FIX --- */
    [data-testid="stTextArea"] textarea {{
        background-color: {theme['input_bg']} !important;
        color: {theme['text_color']} !important;
        border: 1px solid {theme['border_color']} !important;
    }}
    [data-testid="stTextArea"] textarea:focus {{
        border-color: #F39C12 !important;
        box-shadow: 0 0 0 1px #F39C12 !important;
    }}

    /* POPUP DIALOG */
    div[role="dialog"] {{
        background-color: {theme['card_bg']} !important;
        color: {theme['text_color']} !important;
        border: 1px solid {theme['border_color']} !important;
    }}
    div[role="dialog"] h2, div[role="dialog"] p, div[role="dialog"] label {{
        color: {theme['text_color']} !important;
    }}

    /* FILE UPLOADER */
    [data-testid="stFileUploader"] section {{
        background-color: {theme['uploader_bg']} !important;
        border: 1px dashed {theme['border_color']} !important;
    }}
    [data-testid="stFileUploader"] div, 
    [data-testid="stFileUploader"] span, 
    [data-testid="stFileUploader"] small {{
        color: {theme['text_color']} !important;
    }}
    [data-testid="stFileUploader"] button {{
        background-color: {theme['card_bg']} !important;
        color: {theme['text_color']} !important;
        border: 1px solid {theme['border_color']} !important;
    }}

    /* BUTTONS */
    button[kind="primary"] {{
        background: #F39C12 !important; 
        color: white !important;
        border: none;
        border-radius: 8px;
        padding: 12px 20px;
        font-size: 16px;
        font-weight: bold;
        width: 100%; 
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
    }}
    button[kind="primary"]:hover {{
        background: #E67E22 !important; 
        box-shadow: 0 6px 12px rgba(0,0,0,0.2);
        transform: translateY(-2px);
    }}

    div[data-testid="stButton"] button[kind="secondary"] {{
        background-color: transparent !important;
        border: none !important;
        color: {theme['text_color']} !important;
        font-size: 26px;
        padding: 0;
        display: flex; align-items: center; justify-content: center;
        box-shadow: none !important;
    }}
    div[data-testid="stButton"] button[kind="secondary"]:hover {{
        background-color: transparent !important;
        transform: scale(1.2);
        box-shadow: none !important;
    }}
    
    .diff-container {{
        background-color: {theme['diff_bg']};
        color: {theme['diff_text']};
        padding: 25px;
        border-radius: 12px; 
        border: 1px solid {theme['border_color']}; 
        margin-top: 15px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.05);
        font-family: sans-serif;
    }}
    .diff-table {{ width: 100%; border-collapse: collapse; table-layout: fixed; }}
    .diff-table th {{
        text-align: left; padding: 12px; border-bottom: 2px solid {theme['border_color']};
        width: 50%; font-size: 16px; font-weight: 700; color: {theme['text_color']};
    }}
    .diff-cell {{
        width: 50%; padding: 10px; vertical-align: top;
        word-wrap: break-word; border-bottom: 1px solid {theme['border_color']};
        font-size: 14px; line-height: 1.6; white-space: pre-wrap;
        color: {theme['diff_text']};
    }}
    .diff-del {{ background-color: {theme['diff_del_bg']}; color: {theme['diff_del_text']}; text-decoration: line-through; padding: 0 2px; border-radius: 4px; }}
    .diff-add {{ background-color: {theme['diff_add_bg']}; color: {theme['diff_add_text']}; font-weight: 600; padding: 0 2px; border-radius: 4px; }}

    div[data-testid="stVerticalBlock"] > div:has(div#header-marker) {{
        position: sticky;
        top: 2.8rem; 
        background-color: {theme['bg_color']}; 
        z-index: 990; 
        padding-top: 1rem;
        padding-bottom: 1rem;
        border-bottom: 1px solid {theme['border_color']};
        margin-bottom: 1rem;
    }}

    #loading-overlay {{
        position: fixed; top: 0; left: 0; width: 100vw; height: 100vh;
        background-color: {theme['overlay_bg']};
        z-index: 99999;
        display: flex; flex-direction: column; align-items: center; justify-content: center;
        backdrop-filter: blur(5px);
    }}
    
    @keyframes spin {{ 0% {{ transform: rotate(0deg); }} 100% {{ transform: rotate(360deg); }} }}
    .loader-emoji {{ font-size: 80px; animation: spin 1.5s linear infinite; margin-bottom: 20px; }}
    .loader-text {{ font-size: 24px; font-weight: bold; color: #F39C12; font-family: sans-serif; }}
</style>
""", unsafe_allow_html=True)

# --- HEADER SECTION ---
header_container = st.container()
with header_container:
    st.markdown('<div id="header-marker"></div>', unsafe_allow_html=True)
    col_head, col_toggle = st.columns([0.9, 0.1])
    with col_head:
        st.title("✨ Creative Document Difference Checker")
        st.markdown("Upload **Text, PDF, or Word** files below to generate an AI-powered comparison.")
    with col_toggle:
        st.write("") 
        st.write("") 
        btn_emoji = "☀️" if st.session_state.dark_mode else "🌙"
        st.button(btn_emoji, on_click=toggle_theme, key="theme_toggle")

st.write("") 

# --- HELPER FUNCTIONS ---
def read_pdf(file):
    pdf_reader = pypdf.PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() + "\n"
    return text

def read_docx(file):
    doc = docx.Document(file)
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    return "\n".join(text)

def process_uploaded_file(uploaded_file):
    if uploaded_file is None: return ""
    if uploaded_file.name.endswith('.pdf'): return read_pdf(uploaded_file)
    elif uploaded_file.name.endswith('.docx'): return read_docx(uploaded_file)
    else: return uploaded_file.read().decode("utf-8")

def create_docx(text):
    doc = Document()
    doc.add_heading('AI Summary of Changes', 0)
    doc.add_paragraph(text)
    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

@st.dialog("💾 Download Options")
def show_download_dialog(summary_text):
    st.write("Select the format you want to download:")
    format_choice = st.radio("Format", ["Text File (.txt)", "Markdown File (.md)", "Word Document (.docx)"], label_visibility="collapsed")
    
    if "Text" in format_choice:
        file_data, file_name, mime_type = summary_text, "ai_summary.txt", "text/plain"
    elif "Markdown" in format_choice:
        file_data, file_name, mime_type = summary_text, "ai_summary.md", "text/markdown"
    elif "Word" in format_choice:
        file_data, file_name, mime_type = create_docx(summary_text), "ai_summary.docx", "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    
    st.write("") 
    st.download_button(label=f"📥 Confirm Download as {file_name.split('.')[-1].upper()}", data=file_data, file_name=file_name, mime=mime_type, type="primary")

# --- INPUT SECTION ---
col1, col2 = st.columns(2)

with col1:
    with st.container(border=True):
        st.subheader("📄 Document A")
        file_a = st.file_uploader("Upload A", type=['txt', 'pdf', 'docx'], key="ua", label_visibility="collapsed")
        st.caption("Drag and drop .txt, .pdf, or .docx")
        with st.expander("📝 Or paste text manually"):
            text_a_input = st.text_area("Paste Original Text", height=200, key="a")
    text_a = process_uploaded_file(file_a) if file_a else text_a_input

with col2:
    with st.container(border=True):
        st.subheader("📄 Document B")
        file_b = st.file_uploader("Upload B", type=['txt', 'pdf', 'docx'], key="ub", label_visibility="collapsed")
        st.caption("Drag and drop .txt, .pdf, or .docx")
        with st.expander("📝 Or paste text manually"):
            text_b_input = st.text_area("Paste Modified Text", height=200, key="b")
    text_b = process_uploaded_file(file_b) if file_b else text_b_input

# --- TABLE GENERATOR ---
def build_html_table(json_diff):
    rows = ""
    for item in json_diff:
        original = item['original'] or ""
        modified = item['modified'] or ""
        if '<span' not in original: original = html_lib.escape(original)
        if '<span' not in modified: modified = html_lib.escape(modified)
        original = original.replace('\n', '<br>')
        modified = modified.replace('\n', '<br>')
        rows += f'<tr><td class="diff-cell">{original}</td><td class="diff-cell">{modified}</td></tr>'

    html = f"""
    <div class="diff-container">
        <table class="diff-table">
            <thead><tr><th>Original Document</th><th>Modified Document</th></tr></thead>
            <tbody>{rows}</tbody>
        </table>
    </div>
    """
    return html

# --- ACTION SECTION ---
st.write("")
st.write("")
b_col1, b_col2, b_col3 = st.columns([1, 0.6, 1])
with b_col2:
    analyze_clicked = st.button("🚀 Upload & Analyze", type="primary", use_container_width=True)

# --- RESULTS LOGIC ---
if analyze_clicked:
    if not text_a or not text_b:
        st.warning("⚠️ Please upload files or paste text for both documents.")
    else:
        loader_placeholder = st.empty()
        with loader_placeholder.container():
            st.markdown("""<div id="loading-overlay"><div class="loader-emoji">⏳</div><div class="loader-text">Analyzing documents... Please wait</div></div>""", unsafe_allow_html=True)

        try:
            response = requests.post("http://127.0.0.1:8000/compare", json={"text_a": text_a, "text_b": text_b})
            loader_placeholder.empty()

            if response.status_code == 200:
                st.session_state.api_data = response.json()
                st.session_state.analysis_complete = True
            else:
                st.error("Backend Error")
        except Exception as e:
            loader_placeholder.empty()
            st.error(f"Connection failed: {e}")

# --- DISPLAY RESULTS (FROM STATE) ---
if st.session_state.analysis_complete and st.session_state.api_data:
    data = st.session_state.api_data
    
    st.markdown("<div id='results-anchor' style='position: relative; top: -120px; visibility: hidden;'></div>", unsafe_allow_html=True)
    
    st.divider()
    st.subheader("🤖 AI Analysis")
    st.info(data["summary"])
    
    if st.button("📥 Download Summary", type="primary"):
        show_download_dialog(data["summary"])
    
    st.divider()
    st.subheader("📝 Technical Diff")
    if data["diff"].strip():
        st.code(data["diff"], language="diff")
    else:
        st.success("✅ No technical differences found. The documents are identical.")

    st.divider()
    st.subheader("👀 Side-by-Side Comparison")
    html_table = build_html_table(data["json_diff"])
    st.markdown(html_table, unsafe_allow_html=True)
    
    if analyze_clicked:
        components.html("""<script>const anchor = window.parent.document.getElementById('results-anchor'); if (anchor) { anchor.scrollIntoView({behavior: 'smooth', block: 'start'}); }</script>""", height=0)